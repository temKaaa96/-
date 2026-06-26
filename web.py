"""
Веб-версия AI-чата. FastAPI.
Groq (бесплатные модели + vision + распознавание голоса) и Gemini (премиум).

Возможности:
  • текстовый чат со стримингом (Groq / Gemini);
  • фото → разбор изображением (vision): Gemini, либо Groq llama-4-scout;
  • голосовые → распознавание через Groq Whisper (/api/transcribe);
  • файлы (PDF / DOCX / TXT) → извлечение текста (/api/extract);
  • несколько чатов и вложения — на стороне браузера.

Ключи живут на сервере. GEMINI_API_KEY — необязательный ключ владельца.
"""

import os
import io
import json
from pathlib import Path

import httpx
from fastapi import FastAPI, Request, UploadFile, File
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()

GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_WHISPER_URL = "https://api.groq.com/openai/v1/audio/transcriptions"
GEMINI_NATIVE_BASE = os.environ.get(
    "GEMINI_BASE_URL", "https://generativelanguage.googleapis.com/v1beta"
).rstrip("/")

GROQ_VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
GROQ_WHISPER_MODEL = "whisper-large-v3-turbo"

MAX_TOKENS = 2048
MAX_MESSAGES = 24
MAX_CONTENT = 12000

MODELS = {
    "llama-3.3-70b-versatile": {"provider": "groq", "label": "Llama 3.3 70B"},
    "llama-3.1-8b-instant":    {"provider": "groq", "label": "Llama 3.1 8B (быстро)"},
    "gemini-2.5-flash":        {"provider": "gemini", "label": "Gemini 2.5 Flash (премиум)"},
}

app = FastAPI(title="AI Chat")
HERE = Path(__file__).parent

# Папка для картинок оформления (фон, лого и т.п.): кладёшь файлы в static/,
# и они доступны по /static/имя_файла
STATIC_DIR = HERE / "static"
STATIC_DIR.mkdir(exist_ok=True)
# Каталог фонов: кидаешь сюда файлы (mp4/webm/gif/jpg/png) — они появятся в галерее
BG_DIR = STATIC_DIR / "backgrounds"
BG_DIR.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

VIDEO_EXT = {".mp4", ".webm", ".ogv", ".mov"}
IMAGE_EXT = {".jpg", ".jpeg", ".png", ".webp", ".avif"}


class UpstreamError(Exception):
    def __init__(self, status: int, body: str = ""):
        self.status = status
        self.body = body


# ─── Утилиты ─────────────────────────────────────────────────────────────────
def sanitize(messages) -> list:
    out = []
    if isinstance(messages, list):
        for m in messages[-MAX_MESSAGES:]:
            role = m.get("role")
            content = m.get("content")
            if role in ("user", "assistant") and isinstance(content, str) and content.strip():
                out.append({"role": role, "content": content[:MAX_CONTENT]})
    while out and out[0]["role"] != "user":
        out.pop(0)
    return out


def parse_data_url(data_url: str):
    """'data:image/jpeg;base64,XXXX' -> ('image/jpeg', 'XXXX')."""
    try:
        head, b64 = data_url.split(",", 1)
        mime = head.split(":", 1)[1].split(";", 1)[0]
        return mime or "image/jpeg", b64
    except Exception:
        return "image/jpeg", ""


def to_gemini_contents(messages: list, image: str = None) -> list:
    contents = [
        {"role": "model" if m["role"] == "assistant" else "user", "parts": [{"text": m["content"]}]}
        for m in messages
    ]
    if image and contents and contents[-1]["role"] == "user":
        mime, data = parse_data_url(image)
        if data:
            contents[-1]["parts"].insert(0, {"inline_data": {"mime_type": mime, "data": data}})
    return contents


# ─── Стримеры ────────────────────────────────────────────────────────────────
async def groq_stream(model: str, messages: list):
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
    payload = {"model": model, "messages": messages, "max_tokens": MAX_TOKENS, "stream": True}
    async with httpx.AsyncClient(timeout=httpx.Timeout(180.0, connect=20.0)) as c:
        async with c.stream("POST", GROQ_URL, headers=headers, json=payload) as r:
            if r.status_code >= 400:
                raise UpstreamError(r.status_code, (await r.aread()).decode("utf-8", "ignore"))
            async for line in r.aiter_lines():
                line = line.strip()
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if not data or data == "[DONE]":
                    continue
                try:
                    piece = json.loads(data)["choices"][0].get("delta", {}).get("content")
                    if piece:
                        yield piece
                except Exception:
                    continue


async def groq_vision_stream(messages: list, image: str):
    """Последнее сообщение пользователя дополняем картинкой, модель — vision."""
    msgs = [dict(m) for m in messages]
    if msgs and msgs[-1]["role"] == "user":
        msgs[-1] = {
            "role": "user",
            "content": [
                {"type": "text", "text": msgs[-1]["content"]},
                {"type": "image_url", "image_url": {"url": image}},
            ],
        }
    async for piece in groq_stream(GROQ_VISION_MODEL, msgs):
        yield piece


async def gemini_stream(model: str, messages: list, key: str, image: str = None):
    url = f"{GEMINI_NATIVE_BASE}/models/{model}:streamGenerateContent?alt=sse"
    headers = {"x-goog-api-key": key, "Content-Type": "application/json"}
    payload = {
        "contents": to_gemini_contents(messages, image),
        "generationConfig": {"maxOutputTokens": MAX_TOKENS, "thinkingConfig": {"thinkingBudget": 0}},
    }
    async with httpx.AsyncClient(timeout=httpx.Timeout(180.0, connect=20.0)) as c:
        async with c.stream("POST", url, headers=headers, json=payload) as r:
            if r.status_code >= 400:
                raise UpstreamError(r.status_code, (await r.aread()).decode("utf-8", "ignore"))
            async for line in r.aiter_lines():
                line = line.strip()
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if not data:
                    continue
                try:
                    obj = json.loads(data)
                    for part in obj["candidates"][0]["content"]["parts"]:
                        if part.get("text") and not part.get("thought"):
                            yield part["text"]
                except Exception:
                    continue


# ─── Распознавание голоса (Groq Whisper) ─────────────────────────────────────
async def whisper(content: bytes, filename: str) -> str:
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
    files = {"file": (filename, content, "application/octet-stream")}
    data = {"model": GROQ_WHISPER_MODEL, "response_format": "json"}
    async with httpx.AsyncClient(timeout=httpx.Timeout(120.0, connect=20.0)) as c:
        r = await c.post(GROQ_WHISPER_URL, headers=headers, files=files, data=data)
        if r.status_code >= 400:
            raise UpstreamError(r.status_code, r.text)
        return (r.json().get("text") or "").strip()


# ─── Извлечение текста из файла ──────────────────────────────────────────────
def extract_text(content: bytes, name: str, mime: str) -> str:
    name = (name or "").lower()
    if mime == "text/plain" or name.endswith(".txt"):
        try:
            return content.decode("utf-8")[:MAX_CONTENT]
        except UnicodeDecodeError:
            return content.decode("latin-1", "ignore")[:MAX_CONTENT]

    if mime == "application/pdf" or name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text = "\n".join((p.extract_text() or "") for p in reader.pages[:25]).strip()
            if len(text) > 20:
                return text[:MAX_CONTENT]
        except Exception:
            pass
        return "⚠️ Не удалось извлечь текст из PDF (возможно, это скан без текстового слоя)."

    DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if mime == DOCX or name.endswith(".docx"):
        try:
            from docx import Document
            d = Document(io.BytesIO(content))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts).strip()
            if len(text) > 5:
                return text[:MAX_CONTENT]
        except Exception:
            pass
        return "⚠️ Не удалось прочитать DOCX-файл."

    if name.endswith(".doc"):
        return "⚠️ Старый формат .doc не поддерживается. Сохрани файл как .docx."
    return "⚠️ Поддерживаются файлы TXT, PDF, DOCX."


# ─── Эндпоинты ───────────────────────────────────────────────────────────────
@app.get("/")
async def index():
    return FileResponse(HERE / "index.html")


@app.get("/api/models")
async def models():
    return {"models": [{"id": k, "label": v["label"]} for k, v in MODELS.items()]}


@app.get("/api/backgrounds")
async def backgrounds():
    """Список файлов из static/backgrounds/ для галереи фонов."""
    items = []
    if BG_DIR.exists():
        for p in sorted(BG_DIR.iterdir()):
            if not p.is_file():
                continue
            ext = p.suffix.lower()
            if ext in VIDEO_EXT:
                typ = "video"
            elif ext == ".gif":
                typ = "gif"
            elif ext in IMAGE_EXT:
                typ = "image"
            else:
                continue
            items.append({"name": p.stem, "url": f"/static/backgrounds/{p.name}", "type": typ})
    return {"backgrounds": items}


@app.post("/api/transcribe")
async def transcribe(file: UploadFile = File(...)):
    content = await file.read()
    if not content:
        return JSONResponse({"error": "Пустой файл"}, status_code=400)
    try:
        text = await whisper(content, file.filename or "voice.ogg")
    except UpstreamError:
        return JSONResponse({"error": "Не удалось распознать голос."}, status_code=502)
    return {"text": text}


@app.post("/api/extract")
async def extract(file: UploadFile = File(...)):
    content = await file.read()
    if not content:
        return JSONResponse({"error": "Пустой файл"}, status_code=400)
    text = extract_text(content, file.filename or "", file.content_type or "")
    if text.startswith("⚠️"):
        return JSONResponse({"error": text}, status_code=400)
    return {"text": text, "name": file.filename}


@app.post("/api/chat")
async def chat(req: Request):
    body = await req.json()
    model = body.get("model", "llama-3.3-70b-versatile")
    messages = sanitize(body.get("messages", []))
    image = body.get("image")
    gemini_key = (body.get("gemini_key") or body.get("om_key") or GEMINI_API_KEY or "").strip()

    if not messages:
        return JSONResponse({"error": "Пустой запрос"}, status_code=400)

    # выбираем стример
    if image:
        if gemini_key:
            def make(): return gemini_stream("gemini-2.5-flash", messages, gemini_key, image=image)
        else:
            def make(): return groq_vision_stream(messages, image)
    else:
        info = MODELS.get(model)
        if not info:
            return JSONResponse({"error": "Неизвестная модель"}, status_code=400)
        if info["provider"] == "gemini":
            if not gemini_key:
                return JSONResponse(
                    {"error": "Для премиум-модели нужен ключ Gemini — вставь его в настройках (⚙️)."},
                    status_code=400,
                )
            def make(): return gemini_stream(model, messages, gemini_key)
        else:
            def make(): return groq_stream(model, messages)

    async def gen():
        try:
            async for t in make():
                yield t
        except UpstreamError as e:
            if e.status in (401, 403):
                yield "\n\n⚠️ Ключ недействителен."
            elif e.status == 429:
                yield "\n\n⏳ Превышен лимит, подожди минуту."
            else:
                yield f"\n\n⚠️ Ошибка модели ({e.status})."
        except Exception:
            yield "\n\n⚠️ Что-то пошло не так. Попробуй ещё раз."

    return StreamingResponse(gen(), media_type="text/plain; charset=utf-8")
